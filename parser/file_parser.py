import os
import logging
import zipfile
import shutil
from pathlib import Path
from docx import Document
import win32com.client  # 用于处理.doc文件
import pythoncom
from sqlalchemy.orm import Session
from utils.db import get_db, update_project, TenderProject
import time
import platform
import threading
from functools import wraps

# 新增：用于处理rar文件和xlsx文件
import patoolib
from pyunpack import Archive
import openpyxl
import xlrd

# Windows和Unix系统的文件锁模块（可选导入）
try:
    import msvcrt  # Windows文件锁
    MSVCRT_AVAILABLE = True
except ImportError:
    MSVCRT_AVAILABLE = False

try:
    import fcntl  # Unix文件锁（Windows不可用）
    FCNTL_AVAILABLE = True
except ImportError:
    FCNTL_AVAILABLE = False

import psutil

# 可选依赖：PDF和OCR相关（如果未安装，相应功能将不可用）
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

try:
    import pdf2image
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False


class FileParser:
    """文件解析器（修复版）"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = ['pdf', 'docx', 'doc', 'docm', 'txt', 'xlsx', 'xls']  # 支持DOCM格式和Excel文件
        self.archive_formats = ['zip', 'rar']  # 支持zip和rar文件
        # 关键词列表，用于筛选招标文件
        self.tender_keywords = ['招标', '标书', '投标', '采购', '竞争性谈判', '询价', '磋商', '比选', '资格预审']
        # 性能配置
        self.max_file_size_mb = 50  # 最大文件大小（MB），超过此大小会警告
        self.parse_timeout_seconds = 300  # 单个文件解析超时时间（5分钟）
        self.ocr_timeout_seconds = 600  # OCR 解析超时时间（10分钟）
        
        # 检查Word COM组件是否可用（云端环境检测）
        self._word_com_available = self._check_word_com_availability()
        
        # Word COM锁文件路径（用于防止并发访问）
        self._word_lock_file_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                                                 'tender_files', '.word_com_lock')
        self._word_lock_file = None

    def _check_word_com_availability(self):
        """检查Word COM组件是否可用（增强版：带自动恢复和重试）"""
        max_retries = 3
        retry_delay = 2  # 秒
        
        for attempt in range(max_retries):
            try:
                # 先清理可能存在的残留Word进程
                if attempt > 0:
                    self.logger.info(f"尝试恢复Word COM组件（第{attempt + 1}次）...")
                    self._kill_word_processes()
                    time.sleep(retry_delay)
                
                # 初始化COM
                try:
                    pythoncom.CoInitialize()
                except pythoncom.com_error as e:
                    # 如果COM已经初始化，尝试取消初始化后重新初始化
                    if "already initialized" in str(e).lower() or "CoInitialize has not been called" not in str(e):
                        try:
                            pythoncom.CoUninitialize()
                            time.sleep(0.5)
                            pythoncom.CoInitialize()
                        except:
                            pass
                
                # 创建Word应用程序对象
                word = win32com.client.Dispatch("Word.Application")
                if word is None:
                    raise Exception("无法创建Word.Application对象")
                
                word.Visible = False
                word.DisplayAlerts = 0
                
                # 测试是否能正常操作
                doc_count = word.Documents.Count
                
                # 测试成功，清理并返回
                word.Quit(SaveChanges=False)
                time.sleep(0.3)
                pythoncom.CoUninitialize()
                
                # 确保没有残留进程
                self._kill_word_processes()
                
                self.logger.info("✅ Word COM组件检查通过，可用")
                return True
                
            except Exception as e:
                error_msg = str(e)
                error_code_info = ""
                
                # 识别常见错误代码
                if "(-2146959355" in error_msg:
                    error_code_info = "(-2146959355) 服务器运行失败"
                elif "(-2147221021" in error_msg:
                    error_code_info = "(-2147221021) 操作无法使用"
                elif "(-2147024891" in error_msg:
                    error_code_info = "(-2147024891) 拒绝访问（权限不足）"
                elif "(-2147023174" in error_msg:
                    error_code_info = "(-2147023174) RPC服务器不可用"
                else:
                    error_code_info = error_msg[:100]
                
                # 清理资源
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass
                
                # 如果是最后一次尝试，记录详细错误
                if attempt == max_retries - 1:
                    self.logger.warning(f"⚠️ Word COM组件不可用：{error_code_info}")
                    self.logger.info("💡 解决建议：")
                    self.logger.info("   1. 确保已安装Microsoft Word（完整版，非Word Viewer）")
                    self.logger.info("   2. 以管理员身份运行程序，确保有COM组件访问权限")
                    self.logger.info("   3. 检查是否有其他程序占用Word进程")
                    self.logger.info("   4. 尝试重启系统或重新安装Microsoft Office")
                    self.logger.info("   5. 如果无法使用Word，可以将DOC文件手动转换为DOCX格式")
                    return False
                else:
                    # 不是最后一次，继续重试
                    self.logger.debug(f"Word COM组件检查失败（第{attempt + 1}次），将重试：{error_code_info}")
                    continue
        
        return False
    
    def _acquire_word_lock(self, timeout=60):
        """获取Word COM使用锁（防止多进程并发访问冲突）
        
        Args:
            timeout: 获取锁的超时时间（秒），如果超时则返回False
            
        Returns:
            bool: 是否成功获取锁
        """
        try:
            # 确保锁文件目录存在
            lock_dir = os.path.dirname(self._word_lock_file_path)
            os.makedirs(lock_dir, exist_ok=True)
            
            # 尝试获取锁（使用文件锁机制）
            start_time = time.time()
            while time.time() - start_time < timeout:
                try:
                    if platform.system() == 'Windows':
                        # Windows Server使用文件创建作为锁（原子操作，适合Windows Server环境）
                        try:
                            # 尝试以独占模式创建文件（'x'模式在Windows上是原子操作）
                            self._word_lock_file = open(self._word_lock_file_path, 'x')
                            # 写入当前进程ID和时间戳
                            lock_info = f"{os.getpid()}\n{time.time()}"
                            self._word_lock_file.write(lock_info)
                            self._word_lock_file.flush()
                            self.logger.debug(f"成功获取Word COM锁（Windows Server，PID: {os.getpid()}）")
                            return True
                        except FileExistsError:
                            # 文件已存在，检查进程是否还在运行
                            try:
                                with open(self._word_lock_file_path, 'r') as f:
                                    pid_str = f.read().strip()
                                    if pid_str:
                                        pid = int(pid_str)
                                        # 检查进程是否还在运行
                                        if psutil.pid_exists(pid):
                                            # 进程还在运行，等待
                                            time.sleep(0.5)
                                            continue
                                        else:
                                            # 进程已不存在，删除锁文件重试
                                            try:
                                                os.remove(self._word_lock_file_path)
                                            except:
                                                pass
                                            continue
                            except (ValueError, FileNotFoundError, psutil.NoSuchProcess):
                                # 锁文件无效，删除后重试
                                try:
                                    os.remove(self._word_lock_file_path)
                                except:
                                    pass
                                continue
                    else:
                        # Unix/Linux使用fcntl（如果可用）
                        if not FCNTL_AVAILABLE:
                            self.logger.warning("fcntl模块不可用，无法在Unix/Linux系统上使用文件锁")
                            time.sleep(0.5)
                            continue
                        
                        self._word_lock_file = open(self._word_lock_file_path, 'w')
                        try:
                            fcntl.flock(self._word_lock_file.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
                            self._word_lock_file.write(str(os.getpid()))
                            self._word_lock_file.flush()
                            self.logger.debug(f"成功获取Word COM锁（Unix，PID: {os.getpid()}）")
                            return True
                        except IOError:
                            # 锁被占用，等待后重试
                            self._word_lock_file.close()
                            self._word_lock_file = None
                            time.sleep(0.5)
                            continue
                except Exception as e:
                    self.logger.debug(f"获取锁时出错（重试中）：{str(e)}")
                    time.sleep(0.5)
                    continue
            
            # 超时
            self.logger.warning(f"获取Word COM锁超时（{timeout}秒），Word COM可能正被其他进程使用")
            return False
        except Exception as e:
            self.logger.error(f"获取Word COM锁失败：{str(e)}")
            if self._word_lock_file:
                try:
                    self._word_lock_file.close()
                except:
                    pass
                self._word_lock_file = None
            return False
    
    def _release_word_lock(self):
        """释放Word COM使用锁"""
        if self._word_lock_file:
            try:
                if platform.system() == 'Windows':
                    # Windows Server：关闭文件并删除锁文件
                    # 注意：Windows Server上使用文件创建作为锁机制，直接关闭文件即可
                    # 文件会在删除时自动释放，不需要额外的解锁操作
                    self._word_lock_file.close()
                else:
                    # Unix：释放fcntl锁
                    if FCNTL_AVAILABLE:
                        try:
                            fcntl.flock(self._word_lock_file.fileno(), fcntl.LOCK_UN)
                        except:
                            pass
                    self._word_lock_file.close()
                self.logger.debug(f"已释放Word COM锁（PID: {os.getpid()}）")
            except Exception as e:
                self.logger.warning(f"释放Word COM锁时出错：{str(e)}")
            finally:
                self._word_lock_file = None
                # 尝试删除锁文件
                try:
                    if os.path.exists(self._word_lock_file_path):
                        os.remove(self._word_lock_file_path)
                except Exception as e:
                    self.logger.debug(f"删除锁文件失败（可忽略）：{str(e)}")
    
    def _kill_word_processes(self):
        """强制终止所有Word进程（修复堵塞问题）"""
        try:
            killed_count = 0
            for proc in psutil.process_iter(['pid', 'name']):
                try:
                    if proc.info['name'] and 'winword.exe' in proc.info['name'].lower():
                        proc.kill()
                        killed_count += 1
                        self.logger.info(f"已终止Word进程：PID {proc.info['pid']}")
                except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                    pass
            
            if killed_count > 0:
                time.sleep(1)  # 等待进程完全退出
                self.logger.info(f"共清理 {killed_count} 个Word进程")
            return killed_count
        except Exception as e:
            self.logger.warning(f"清理Word进程时出错：{str(e)}")
            return 0

    def _is_tender_file(self, file_name):
        """判断文件是否为招标文件（改进版：更宽松的识别）"""
        file_name_lower = file_name.lower()
        
        # 检查是否包含关键词
        for keyword in self.tender_keywords:
            if keyword in file_name_lower:
                return True
        
        # 如果文件名包含"项目"和"文件"，也可能是招标文件
        if '项目' in file_name_lower and '文件' in file_name_lower:
            return True
        
        return False

    def _extract_zip(self, archive_path):
        """解压压缩文件（zip或rar）并返回相关文件路径列表（改进版：更好的错误处理和文件识别）"""
        try:
            extract_dir = os.path.splitext(archive_path)[0]  # 使用压缩文件同名目录解压
            os.makedirs(extract_dir, exist_ok=True)
            
            file_ext = Path(archive_path).suffix.lower().lstrip('.')
            all_files = []
            
            if file_ext == 'zip':
                # 处理zip文件
                with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                    all_files = zip_ref.namelist()
                    
                    for file in all_files:
                        file_name = os.path.basename(file)
                        # 跳过隐藏文件和目录
                        if file_name.startswith('.') or file.endswith('/'):
                            continue
                        
                        # 解压文件
                        try:
                            zip_ref.extract(file, extract_dir)
                        except Exception as e:
                            self.logger.error(f"解压zip文件失败 {file}: {str(e)}")
                            continue
            elif file_ext == 'rar':
                # 处理rar文件
                try:
                    from pyunpack import Archive
                    Archive(archive_path).extractall(extract_dir)
                    # 获取解压后的所有文件
                    for root, dirs, files in os.walk(extract_dir):
                        for file in files:
                            rel_path = os.path.relpath(os.path.join(root, file), extract_dir)
                            all_files.append(rel_path)
                except Exception as e:
                    self.logger.error(f"解压rar文件失败 {archive_path}：{str(e)}")
                    import traceback
                    self.logger.error(traceback.format_exc())
                    return []
            
            # 收集招标文件
            tender_files = []
            supported_files = []  # 支持格式的文件（即使不包含关键词）
            
            for file in all_files:
                file_name = os.path.basename(file)
                # 跳过隐藏文件和目录
                if file_name.startswith('.') or file.endswith('/'):
                    continue
                
                # 处理路径：如果压缩文件中有子目录，需要正确拼接路径
                extracted_path = os.path.join(extract_dir, file)
                # 标准化路径，处理Windows路径分隔符问题
                extracted_path = os.path.normpath(extracted_path)
                
                # 确保文件存在
                if not os.path.exists(extracted_path):
                    # 尝试使用绝对路径
                    extracted_path = os.path.abspath(extracted_path)
                    if not os.path.exists(extracted_path):
                        self.logger.warning(f"解压后文件不存在: {extracted_path}")
                        # 尝试查找文件（可能路径编码问题）
                        file_name_only = os.path.basename(file)
                        for root, dirs, files in os.walk(extract_dir):
                            if file_name_only in files:
                                extracted_path = os.path.join(root, file_name_only)
                                self.logger.info(f"找到文件（使用搜索）: {extracted_path}")
                                break
                        else:
                            continue
                
                # 判断是否为招标文件
                if self._is_tender_file(file_name):
                    tender_files.append(extracted_path)
                    self.logger.info(f"从{file_ext}中提取并识别为招标文件：{extracted_path}")
                else:
                    # 即使不包含关键词，如果是支持的格式，也记录
                    file_ext_inner = os.path.splitext(file_name)[1].lower().lstrip('.')
                    if file_ext_inner in self.supported_formats:
                        supported_files.append(extracted_path)
                        self.logger.info(f"从{file_ext}中提取文件（格式支持）：{extracted_path}")
                    else:
                        self.logger.info(f"从{file_ext}中提取但不参与分析：{extracted_path}")
            
            # 如果没有找到明确的招标文件，但压缩文件中有支持格式的文件，返回所有支持格式的文件
            if not tender_files and supported_files:
                self.logger.warning(f"{file_ext}中未找到明确的招标文件，尝试解析所有支持格式的文件（{len(supported_files)}个）")
                return supported_files
            
            return tender_files
        except Exception as e:
            self.logger.error(f"解压压缩文件失败 {archive_path}：{str(e)}")
            import traceback
            self.logger.error(traceback.format_exc())
            return []

    def parse_file(self, file_path, project_id):
        """解析单个文件（增加错误处理和zip解压支持）"""
        self.logger.info(f"========== 开始解析文件 ==========")
        self.logger.info(f"文件路径: {file_path}")
        self.logger.info(f"项目ID: {project_id}")
        try:
            # 1. 检查文件是否存在
            if not os.path.exists(file_path):
                self.logger.error(f"文件不存在：{file_path}")
                return None

            # 2. 获取文件扩展名
            file_ext = Path(file_path).suffix.lower().lstrip('.')

            # 3. 处理压缩文件
            if file_ext in self.archive_formats:
                self.logger.info(f"处理压缩文件：{file_path}")
                # 解压zip文件并获取招标文件列表
                tender_files = self._extract_zip(file_path)
                
                if not tender_files:
                    self.logger.warning(f"zip文件中未找到招标文件：{file_path}")
                    # 检查解压目录中是否有文件（可能已经解压过，但ZIP文件被删除了）
                    extract_dir = os.path.splitext(file_path)[0]
                    if os.path.exists(extract_dir):
                        self.logger.info(f"ZIP文件可能已解压，检查解压目录: {extract_dir}")
                        for root, dirs, files in os.walk(extract_dir):
                            for file in files:
                                file_path_full = os.path.join(root, file)
                                file_ext_inner = os.path.splitext(file)[1].lower().lstrip('.')
                                if file_ext_inner in self.supported_formats:
                                    if self._is_tender_file(file):
                                        tender_files.append(file_path_full)
                                        self.logger.info(f"在已解压目录中找到招标文件: {file_path_full}")
                    
                    if not tender_files:
                        self.logger.error(f"zip文件中未找到可解析的招标文件：{file_path}")
                        return None
                
                # 解析所有识别的招标文件
                all_content = []
                for tender_file in tender_files:
                    self.logger.info(f"开始解析ZIP中的文件: {tender_file}")
                    # 确保文件存在
                    if not os.path.exists(tender_file):
                        self.logger.error(f"ZIP中的文件不存在: {tender_file}")
                        continue
                    # 检查文件大小
                    if os.path.getsize(tender_file) < 100:
                        self.logger.warning(f"ZIP中的文件过小（{os.path.getsize(tender_file)}字节）: {tender_file}")
                        continue
                    file_content = self.parse_file(tender_file, project_id)
                    if file_content:
                        # 检查解析内容是否为空
                        if file_content.strip():
                            all_content.append(file_content)
                            self.logger.info(f"ZIP文件解析成功: {tender_file}, 内容长度: {len(file_content)} 字符")
                        else:
                            self.logger.warning(f"ZIP文件解析后内容为空: {tender_file}")
                    else:
                        self.logger.warning(f"ZIP文件解析失败（返回None）: {tender_file}")
                
                if all_content:
                    # 解析成功后删除原zip文件，只保留解压后的文件夹
                    try:
                        os.remove(file_path)
                        self.logger.info(f"解析完成后删除原zip文件：{file_path}")
                        
                        # 更新数据库中的file_path字段
                        if project_id:
                            try:
                                # 获取数据库会话
                                db: Session = next(get_db())
                                
                                # 确定新的文件路径
                                if tender_files:
                                    if len(tender_files) == 1:
                                        # 如果只有一个招标文件，直接指向该文件
                                        new_file_path = tender_files[0]
                                    else:
                                        # 如果有多个招标文件，指向解压目录
                                        new_file_path = os.path.splitext(file_path)[0]
                                else:
                                    # 如果没有找到招标文件，指向解压目录
                                    new_file_path = os.path.splitext(file_path)[0]
                                
                                # 更新数据库
                                update_project(db, project_id, {"file_path": new_file_path})
                                self.logger.info(f"更新项目文件路径：{project_id} -> {new_file_path}")
                            except Exception as db_e:
                                self.logger.error(f"更新项目文件路径失败：{project_id}，错误：{str(db_e)}")
                    except Exception as e:
                        self.logger.error(f"删除zip文件失败 {file_path}：{str(e)}")
                    return '\n\n--- 分割线：来自多个文件的内容 ---\n\n'.join(all_content)
                else:
                    self.logger.warning(f"zip文件中的招标文件解析失败：{file_path}")
                    return None
            
            # 4. 处理普通文件格式
            if file_ext not in self.supported_formats:
                self.logger.warning(f"不支持的文件格式：{file_ext}")
                return None

            # 5. 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size < 100:
                self.logger.warning(f"文件过小（{file_size}字节），可能为空或损坏: {file_path}")
                # 对于小文件，尝试解析，但添加警告

            # 6. 根据格式解析
            if file_ext == 'pdf':
                result = self._parse_pdf(file_path)
            elif file_ext == 'docx' or file_ext == 'docm':
                # DOCM是启用宏的Word文档，尝试用_parse_docx解析（内部会处理）
                result = self._parse_docx(file_path)
            elif file_ext == 'doc':
                result = self._parse_doc(file_path)
            elif file_ext == 'txt':
                result = self._parse_txt(file_path)
            elif file_ext == 'xlsx' or file_ext == 'xls':
                result = self._parse_excel(file_path)
            else:
                self.logger.warning(f"不支持的文件格式：{file_ext}")
                result = None

            # 7. 检查解析结果
            if result:
                if result.strip():
                    self.logger.info(f"文件解析成功，内容长度: {len(result)} 字符")
                    return result
                else:
                    self.logger.warning(f"文件解析后内容为空: {file_path}")
                    return None
            else:
                self.logger.warning(f"文件解析失败（返回None）: {file_path}")
                return None

        except Exception as e:
            self.logger.error(f"解析文件失败 {file_path}：{str(e)}", exc_info=True)
            import traceback
            self.logger.error(f"详细错误信息：{traceback.format_exc()}")
            return None
        finally:
            self.logger.info(f"========== 文件解析结束 ==========")

    def _parse_doc_with_libreoffice(self, file_path):
        """使用LibreOffice命令行工具将DOC转换为DOCX，然后解析（备用方案）"""
        try:
            import subprocess
            import tempfile
            import shutil
            
            # 检查LibreOffice是否可用
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                "soffice",  # 如果在PATH中
            ]
            
            soffice_exe = None
            for path in libreoffice_paths:
                if path == "soffice":
                    # 检查是否在PATH中
                    try:
                        result = subprocess.run(
                            ["soffice", "--version"], 
                            capture_output=True, 
                            timeout=5,
                            shell=False,  # 明确指定不使用shell
                            encoding='utf-8',
                            errors='replace'
                        )
                        if result.returncode == 0:
                            soffice_exe = "soffice"
                            break
                    except:
                        continue
                elif os.path.exists(path):
                    soffice_exe = path
                    break
            
            if not soffice_exe:
                self.logger.debug("LibreOffice未找到，无法使用备用方案")
                return None
            
            self.logger.info(f"尝试使用LibreOffice转换DOC文件：{os.path.basename(file_path)}")
            
            # 创建临时目录用于转换
            temp_dir = tempfile.mkdtemp()
            try:
                # 确保使用绝对路径，避免路径中包含特殊字符（如方括号）导致的问题
                abs_file_path = os.path.abspath(file_path)
                abs_temp_dir = os.path.abspath(temp_dir)
                
                # 使用LibreOffice将DOC转换为DOCX
                # --headless: 无界面模式
                # --convert-to docx: 转换为DOCX格式
                # --outdir: 输出目录
                # 注意：使用列表形式传递参数，避免shell解析特殊字符
                cmd = [
                    soffice_exe,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", abs_temp_dir,
                    abs_file_path
                ]
                
                # 在Windows上明确指定shell=False，避免PowerShell解析特殊字符
                result = subprocess.run(
                    cmd, 
                    capture_output=True, 
                    timeout=60, 
                    text=True,
                    shell=False,  # 明确指定不使用shell，避免特殊字符解析问题
                    encoding='utf-8',  # 指定编码
                    errors='replace'  # 遇到编码错误时替换而不是失败
                )
                
                if result.returncode != 0:
                    self.logger.warning(f"LibreOffice转换失败：{result.stderr}")
                    return None
                
                # 查找转换后的文件（LibreOffice会生成同名的docx文件）
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                converted_file = os.path.join(temp_dir, f"{base_name}.docx")
                
                if not os.path.exists(converted_file):
                    self.logger.warning(f"LibreOffice转换后的文件不存在：{converted_file}")
                    return None
                
                # 使用python-docx解析转换后的DOCX文件
                self.logger.info("使用python-docx解析转换后的DOCX文件")
                return self._parse_docx(converted_file)
                
            finally:
                # 清理临时目录
                try:
                    shutil.rmtree(temp_dir, ignore_errors=True)
                except:
                    pass
                    
        except subprocess.TimeoutExpired:
            self.logger.warning("LibreOffice转换超时")
            return None
        except Exception as e:
            self.logger.warning(f"LibreOffice备用方案失败：{str(e)}")
            return None

    def _parse_doc(self, file_path):
        """解析.doc文件（增强版：添加超时和进程清理，云端环境兼容，支持大文件，带备用方案）"""
        start_time = time.time()
        word = None
        doc = None
        
        try:
            # === 云端环境检查：如果Word COM不可用，尝试重新检测 ===
            if not self._word_com_available:
                # 尝试重新检测Word COM组件（可能之前检测失败但现在可用）
                self.logger.info("Word COM组件之前检测为不可用，尝试重新检测...")
                self._word_com_available = self._check_word_com_availability()
                
                if not self._word_com_available:
                    # 尝试使用LibreOffice备用方案
                    self.logger.info("Word COM组件不可用，尝试使用LibreOffice备用方案...")
                    result = self._parse_doc_with_libreoffice(file_path)
                    if result:
                        return result
                    
                    # 使用WARNING级别，因为这是已知的、预期的错误（已在初始化时记录过详细信息）
                    self.logger.warning(f"⚠️ Word COM组件不可用，且LibreOffice备用方案也失败，跳过DOC文件解析：{os.path.basename(file_path)}")
                    # 返回None，让调用者知道解析失败
                    return None
                else:
                    self.logger.info("✅ Word COM组件重新检测成功，可以解析DOC文件")
            
            # === 关键修复：获取Word COM锁，防止并发访问冲突 ===
            if not self._acquire_word_lock(timeout=60):
                self.logger.warning("Word COM组件正被其他进程使用，无法解析DOC文件（请稍后重试或等待其他解析完成）")
                return None
            
            # === 关键修复1：解析前清理残留进程 ===
            self._kill_word_processes()
            time.sleep(0.5)  # 等待清理完成
            
            # 检查文件大小
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            
            # 对于大文件，调整超时时间
            if file_size_mb > 2:
                timeout = min(self.parse_timeout_seconds * 2, 600)  # 最多10分钟
                self.logger.info(f"检测到大文件（{file_size_mb:.2f}MB），使用扩展超时时间：{timeout}秒")
            else:
                timeout = self.parse_timeout_seconds
            
            if file_size_mb > self.max_file_size_mb:
                self.logger.warning(f"DOC文件较大（{file_size_mb:.2f}MB），解析可能需要较长时间：{file_path}")
                self.logger.info(f"提示：大文件可能包含大量图片或嵌入对象，解析时间会较长")
            
            # 使用绝对路径
            abs_path = os.path.abspath(file_path)
            self.logger.info(f"开始解析DOC文件：{abs_path}（大小：{file_size_mb:.2f}MB）")

            # 初始化COM对象并创建Word应用程序对象（带重试机制）
            max_retries = 3
            word = None
            com_initialized = False
            
            for attempt in range(max_retries):
                try:
                    # 初始化COM
                    try:
                        pythoncom.CoInitialize()
                        com_initialized = True
                    except pythoncom.com_error as e:
                        error_str = str(e).lower()
                        # 如果COM已经初始化，这是正常的
                        if "already initialized" in error_str or "coinitialize has not been called" not in error_str:
                            com_initialized = True
                        elif attempt < max_retries - 1:
                            self.logger.warning(f"COM初始化失败（第{attempt + 1}次），尝试恢复...")
                            self._kill_word_processes()
                            time.sleep(1)
                            try:
                                pythoncom.CoUninitialize()
                            except:
                                pass
                            continue
                        else:
                            self.logger.error(f"COM初始化失败：{str(e)}，可能是权限问题或环境配置问题")
                            self._release_word_lock()
                            return None
                    except Exception as e:
                        if attempt < max_retries - 1:
                            self.logger.warning(f"COM初始化失败（第{attempt + 1}次），尝试恢复...")
                            self._kill_word_processes()
                            time.sleep(1)
                            try:
                                pythoncom.CoUninitialize()
                            except:
                                pass
                            continue
                        else:
                            self.logger.error(f"COM初始化失败：{str(e)}，可能是权限问题或环境配置问题")
                            self._release_word_lock()
                            return None
                    
                    # 创建Word应用程序对象
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        if word is not None:
                            break  # 成功创建，退出重试循环
                    except Exception as dispatch_error:
                        error_str = str(dispatch_error)
                        error_code = None
                        
                        # 尝试提取错误代码
                        try:
                            import pywintypes
                            if isinstance(dispatch_error, pywintypes.com_error):
                                error_code = dispatch_error.args[0] if dispatch_error.args else None
                        except (ImportError, AttributeError):
                            pass
                        
                        # 检查是否是COM组件错误
                        is_com_error = (
                            error_code in [-2146959355, -2147221021, -2147023170] or
                            "服务器运行失败" in error_str or
                            "操作无法使用" in error_str or
                            "远程过程调用失败" in error_str
                        )
                        
                        if is_com_error:
                            if attempt < max_retries - 1:
                                error_detail = f"错误代码：{error_code}" if error_code else f"错误信息：{error_str}"
                                self.logger.warning(f"Word COM组件初始化失败（第{attempt + 1}次），尝试恢复：{error_detail}")
                                self._kill_word_processes()
                                time.sleep(1)
                                try:
                                    pythoncom.CoUninitialize()
                                except:
                                    pass
                                continue
                            else:
                                error_detail = f"错误代码：{error_code}" if error_code else f"错误信息：{error_str}"
                                self.logger.error(f"Word COM组件初始化失败：{error_detail}")
                                self.logger.warning("建议：1) 检查Microsoft Word是否正常运行 2) 重启Word应用程序 3) 检查COM组件权限 4) 将DOC文件手动转换为DOCX格式")
                                self._release_word_lock()
                                return None
                        else:
                            # 其他错误，如果是最后一次尝试则抛出
                            if attempt < max_retries - 1:
                                self.logger.warning(f"创建Word.Application失败（第{attempt + 1}次），尝试恢复：{error_str}")
                                self._kill_word_processes()
                                time.sleep(1)
                                try:
                                    pythoncom.CoUninitialize()
                                except:
                                    pass
                                continue
                            else:
                                raise
                
                except Exception as e:
                    if attempt < max_retries - 1:
                        self.logger.warning(f"Word COM初始化失败（第{attempt + 1}次），尝试恢复：{str(e)}")
                        self._kill_word_processes()
                        time.sleep(1)
                        try:
                            pythoncom.CoUninitialize()
                        except:
                            pass
                        continue
                    else:
                        self.logger.error(f"Word COM初始化最终失败：{str(e)}")
                        self._release_word_lock()
                        return None
            
            if word is None:
                self.logger.error("无法创建Word.Application对象，可能是Word未安装或权限不足")
                self._release_word_lock()
                return None
            
            word.Visible = False
            word.DisplayAlerts = 0
            
            # 设置超时：打开文档
            open_start = time.time()
            # === 关键修复2：使用ReadOnly模式打开 ===
            try:
                doc = word.Documents.Open(abs_path, ReadOnly=True)
            except Exception as open_error:
                    error_str = str(open_error)
                    error_code = None
                    
                    # 尝试提取错误代码
                    try:
                        import pywintypes
                        if isinstance(open_error, pywintypes.com_error):
                            error_code = open_error.args[0]
                    except (ImportError, AttributeError):
                        pass
                    
                    # 检查是否是RPC错误（-2147023170：远程过程调用失败）
                    is_rpc_error = (
                        error_code == -2147023170 or
                        '-2147023170' in error_str or
                        '远程过程调用失败' in error_str or
                        'RPC' in error_str.upper()
                    )
                    
                    if is_rpc_error:
                        self.logger.error(f"Word COM RPC错误（-2147023170），可能是Word进程异常或COM组件问题：{abs_path}")
                        self.logger.info("尝试清理Word进程并使用LibreOffice备用方案...")
                        
                        # 清理Word进程
                        try:
                            self._kill_word_processes()
                            time.sleep(1)
                        except:
                            pass
                        
                        # 确保释放锁
                        if self._word_lock_file:
                            self._release_word_lock()
                        
                        # 尝试使用LibreOffice备用方案
                        result = self._parse_doc_with_libreoffice(file_path)
                        if result:
                            self.logger.info("✅ LibreOffice备用方案成功解析DOC文件（RPC错误后）")
                            return result
                        else:
                            self.logger.error("LibreOffice备用方案也失败，无法解析DOC文件")
                            return None
                    
                    # 检查是否是Office检测到文件问题的错误（错误代码 -2147352567 或 -2146821993）
                    is_file_problem = (
                        error_code == -2147352567 or 
                        error_code == -2146821993 or
                        '-2147352567' in error_str or 
                        '-2146821993' in error_str or
                        'Office 检测到此文件存在一个问题' in error_str or 
                        '不能打开此文件' in error_str or
                        '为帮助保护您的计算机' in error_str
                    )
                    
                    if is_file_problem:
                        self.logger.error(f"Word检测到文件存在问题，无法打开（文件可能损坏或包含恶意内容）：{abs_path}")
                        self.logger.error(f"错误代码：{error_code}，错误详情：{error_str}")
                        # 尝试使用备用方法：直接读取文件内容（如果可能）
                        try:
                            # 尝试使用文本模式读取（可能只能读取部分内容）
                            with open(abs_path, 'rb') as f:
                                # 检查文件是否真的是Word文档
                                header = f.read(8)
                                if header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':  # OLE2文件头
                                    self.logger.warning(f"文件是有效的OLE2格式，但Word拒绝打开，可能是文件损坏或安全限制")
                                else:
                                    self.logger.warning(f"文件格式可能不是标准的Word文档")
                        except Exception as read_error:
                            self.logger.debug(f"尝试读取文件头失败：{str(read_error)}")
                        
                        # 尝试使用LibreOffice备用方案
                        if self._word_lock_file:
                            self._release_word_lock()
                        result = self._parse_doc_with_libreoffice(file_path)
                        if result:
                            self.logger.info("✅ LibreOffice备用方案成功解析DOC文件（文件问题后）")
                            return result
                        
                        # 返回None，让调用者知道文件无法解析
                        return None
                    else:
                        self.logger.error(f"打开Word文档失败：{error_str}，文件路径：{abs_path}")
                        # 对于其他错误，也尝试使用LibreOffice备用方案
                        if self._word_lock_file:
                            self._release_word_lock()
                        result = self._parse_doc_with_libreoffice(file_path)
                        if result:
                            self.logger.info("✅ LibreOffice备用方案成功解析DOC文件（其他错误后）")
                            return result
                        raise
            
            if doc is None:
                raise Exception("文档对象为None，可能是文件损坏或无法访问")
            
            if time.time() - open_start > 60:  # 打开超过1分钟警告
                self.logger.warning(f"DOC文件打开耗时较长：{time.time() - open_start:.2f}秒")

            # === 关键优化：提取段落和表格内容 ===
            full_text_parts = []
            
            # 1. 提取段落文本
            try:
                paragraphs = []
                for para in doc.Paragraphs:
                    para_text = para.Range.Text.strip()
                    if para_text:
                        paragraphs.append(para_text)
                if paragraphs:
                    full_text_parts.append('\n'.join(paragraphs))
                    self.logger.debug(f"提取了 {len(paragraphs)} 个段落")
            except Exception as e:
                self.logger.warning(f"提取段落失败：{str(e)}，尝试使用Content.Text")
                try:
                    text_content = doc.Content.Text
                    if text_content and text_content.strip():
                        full_text_parts.append(text_content.strip())
                except Exception as e2:
                    self.logger.warning(f"使用Content.Text也失败：{str(e2)}")
            
            # 2. 提取表格内容（关键优化：确保评分表不丢失，支持合并单元格）
            try:
                table_count = doc.Tables.Count
                self.logger.info(f"检测到 {table_count} 个表格，开始提取表格内容")
                
                for table_idx in range(1, table_count + 1):
                        try:
                            table = doc.Tables(table_idx)
                            # 添加表格标识
                            full_text_parts.append("\n[表格开始]")
                            
                            # === 修复：处理有合并单元格的表格 ===
                            # 方法1：尝试通过单元格索引访问（更健壮，支持合并单元格）
                            try:
                                # 获取表格的行数和列数
                                row_count = table.Rows.Count
                                col_count = table.Columns.Count
                                
                                # 通过行列索引访问单元格（避免合并单元格问题）
                                table_rows = []
                                for row_idx in range(1, row_count + 1):
                                    row_cells = []
                                    for col_idx in range(1, col_count + 1):
                                        try:
                                            # 直接通过行列索引访问单元格（支持合并单元格）
                                            cell = table.Cell(row_idx, col_idx)
                                            cell_text = cell.Range.Text.strip()
                                            # 移除Word表格单元格末尾的换行符和制表符
                                            cell_text = cell_text.replace('\r', '').replace('\n', ' ').replace('\t', ' ').strip()
                                            row_cells.append(cell_text)
                                        except Exception as cell_e:
                                            # 如果单个单元格访问失败，使用空字符串
                                            row_cells.append("")
                                    
                                    # 只添加非空行
                                    if any(cell.strip() for cell in row_cells):
                                        table_rows.append("\t".join(row_cells))
                                
                                # 添加提取的行
                                full_text_parts.extend(table_rows)
                                self.logger.debug(f"表格 {table_idx} 提取完成（方法1：单元格索引），共 {len(table_rows)} 行")
                                
                            except Exception as method1_error:
                                # 方法1失败，尝试方法2：直接获取表格范围文本
                                self.logger.warning(f"表格{table_idx}方法1失败（{str(method1_error)[:100]}），尝试方法2")
                                try:
                                    # 方法2：直接获取整个表格的文本（简单但可能格式不完美）
                                    table_text = table.Range.Text.strip()
                                    if table_text:
                                        # 将表格文本按行分割，用制表符分隔（如果可能）
                                        table_lines = [line.strip() for line in table_text.split('\r') if line.strip()]
                                        full_text_parts.extend(table_lines)
                                        self.logger.debug(f"表格 {table_idx} 提取完成（方法2：范围文本），共 {len(table_lines)} 行")
                                    else:
                                        self.logger.warning(f"表格{table_idx}方法2提取的文本为空")
                                except Exception as method2_error:
                                    # 方法2也失败，尝试方法3：逐行访问（原始方法，但添加更多错误处理）
                                    self.logger.warning(f"表格{table_idx}方法2失败（{str(method2_error)[:100]}），尝试方法3")
                                    try:
                                        table_rows = []
                                        for row_idx in range(1, table.Rows.Count + 1):
                                            try:
                                                row = table.Rows(row_idx)
                                                row_cells = []
                                                # 尝试获取行的单元格数量
                                                try:
                                                    cell_count = row.Cells.Count
                                                    for cell_idx in range(1, cell_count + 1):
                                                        try:
                                                            cell_text = row.Cells(cell_idx).Range.Text.strip()
                                                            cell_text = cell_text.replace('\r', '').replace('\n', ' ').replace('\t', ' ').strip()
                                                            row_cells.append(cell_text)
                                                        except:
                                                            row_cells.append("")
                                                except:
                                                    # 如果无法获取单元格数量，跳过这一行
                                                    continue
                                                
                                                if any(cell.strip() for cell in row_cells):
                                                    table_rows.append("\t".join(row_cells))
                                            except Exception as row_e:
                                                # 如果某行访问失败（可能是合并单元格），跳过该行
                                                self.logger.debug(f"表格{table_idx}行{row_idx}访问失败（可能是有合并单元格），跳过：{str(row_e)[:50]}")
                                                continue
                                        
                                        if table_rows:
                                            full_text_parts.extend(table_rows)
                                            self.logger.debug(f"表格 {table_idx} 提取完成（方法3：逐行访问），共 {len(table_rows)} 行")
                                        else:
                                            self.logger.warning(f"表格{table_idx}方法3未提取到任何行")
                                    except Exception as method3_error:
                                        self.logger.error(f"表格{table_idx}所有方法都失败，最后错误：{str(method3_error)[:100]}")
                            
                            # 添加表格结束标识
                            full_text_parts.append("[表格结束]\n")
                            
                        except Exception as table_e:
                            error_msg = str(table_e)
                            if "纵向合并的单元格" in error_msg or "合并" in error_msg:
                                self.logger.warning(f"提取表格{table_idx}失败（合并单元格问题）：{error_msg[:100]}")
                                # 对于合并单元格问题，尝试使用备用方法
                                try:
                                    table = doc.Tables(table_idx)
                                    full_text_parts.append("\n[表格开始]")
                                    # 直接获取表格范围文本
                                    table_text = table.Range.Text.strip()
                                    if table_text:
                                        table_lines = [line.strip() for line in table_text.split('\r') if line.strip()]
                                        full_text_parts.extend(table_lines)
                                        full_text_parts.append("[表格结束]\n")
                                        self.logger.info(f"表格{table_idx}使用备用方法提取成功，共 {len(table_lines)} 行")
                                    else:
                                        self.logger.warning(f"表格{table_idx}备用方法提取的文本为空")
                                except Exception as backup_error:
                                    self.logger.error(f"表格{table_idx}备用方法也失败：{str(backup_error)[:100]}")
                            else:
                                self.logger.warning(f"提取表格{table_idx}失败：{error_msg[:100]}")
                            continue
                
                if table_count > 0:
                    self.logger.info(f"成功提取了 {table_count} 个表格")
            except Exception as table_err:
                self.logger.warning(f"提取表格时出错：{str(table_err)}，继续使用文本内容")
            
            # 合并所有内容
            text = '\n'.join(full_text_parts) if full_text_parts else ""
            
            # 如果没有提取到任何内容，尝试使用Content.Text作为备用方案
            if not text or not text.strip():
                self.logger.warning(f"DOC文件解析后内容为空，尝试使用Content.Text作为备用方案：{file_path}")
                try:
                    # 安全地访问Content.Text，避免win32com属性访问错误
                    if doc and hasattr(doc, 'Content'):
                        content_obj = doc.Content
                        if content_obj and hasattr(content_obj, 'Text'):
                            text = content_obj.Text
                            self.logger.info(f"使用Content.Text获取内容，长度：{len(text) if text else 0}")
                        else:
                            self.logger.warning("doc.Content对象不存在或没有Text属性")
                    else:
                        self.logger.warning("doc对象不存在或没有Content属性")
                except AttributeError as attr_error:
                    self.logger.warning(f"访问Content.Text时属性错误：{str(attr_error)}")
                except Exception as e:
                    self.logger.warning(f"使用Content.Text失败：{str(e)}")

            # === 关键修复3：立即关闭文档 ===
            try:
                if doc:
                    doc.Close(SaveChanges=False)
            except AttributeError as attr_error:
                # 处理win32com属性访问错误（如Open.Close）
                self.logger.warning(f"关闭文档时属性错误（可忽略）：{str(attr_error)}")
            except Exception as close_error:
                self.logger.warning(f"关闭文档时出错（可忽略）：{str(close_error)}")
            finally:
                doc = None
            
            elapsed = time.time() - start_time
            text_length = len(text) if text else 0
            self.logger.info(f"DOC文件解析完成，耗时：{elapsed:.2f}秒，文本长度：{text_length}字符")
            
            if not text or not text.strip():
                self.logger.error(f"DOC文件解析后内容仍为空：{file_path}，文件大小：{os.path.getsize(file_path)}字节")
                self.logger.error("可能原因：1) Word COM组件权限不足 2) 文件损坏 3) 云端环境配置问题")
                # 尝试使用LibreOffice备用方案
                self.logger.info("尝试使用LibreOffice备用方案解析空内容DOC文件...")
                result = self._parse_doc_with_libreoffice(file_path)
                if result:
                    self.logger.info("✅ LibreOffice备用方案成功解析DOC文件")
                    # 确保释放锁
                    if self._word_lock_file:
                        self._release_word_lock()
                    return result
                # 如果备用方案也失败，返回None
                return None
            
            return text.strip()

        except Exception as e:
            self.logger.error(f"DOC文件解析过程中出错：{str(e)}", exc_info=True)
            # 确保文档关闭
            if doc:
                try:
                    doc.Close(SaveChanges=False)
                except:
                    pass
                doc = None
            # 如果是Word COM错误，直接返回None（不再尝试转换，因为转换也需要Word COM）
            error_str = str(e)
            error_code = None
            
            # 尝试提取错误代码
            try:
                import pywintypes
                if isinstance(e, pywintypes.com_error):
                    error_code = e.args[0] if e.args else None
            except (ImportError, AttributeError):
                pass
            
            # 检查是否是Word COM错误
            is_word_com_error = (
                error_code in [-2146959355, -2147221021, -2147023170] or
                "服务器运行失败" in error_str or
                "操作无法使用" in error_str or
                "远程过程调用失败" in error_str or
                "Word" in error_str or 
                "COM" in error_str or 
                "Application" in error_str
            )
            
            if is_word_com_error:
                error_detail = f"错误代码：{error_code}" if error_code else f"错误信息：{error_str}"
                self.logger.error(f"Word COM组件出错，无法解析DOC文件。{error_detail}")
                
                # 尝试使用LibreOffice备用方案
                self.logger.info("尝试使用LibreOffice备用方案解析DOC文件...")
                result = self._parse_doc_with_libreoffice(file_path)
                if result:
                    self.logger.info("✅ LibreOffice备用方案成功解析DOC文件")
                    # 确保释放锁
                    if self._word_lock_file:
                        self._release_word_lock()
                    return result
                
                self.logger.warning("建议：1) 检查Microsoft Word是否正常运行 2) 重启Word应用程序 3) 检查COM组件权限 4) 安装LibreOffice以使用备用解析方案 5) 将DOC文件手动转换为DOCX格式")
                
                # 尝试清理Word进程
                try:
                    self._kill_word_processes()
                    time.sleep(1)  # 等待清理完成
                except Exception as kill_error:
                    self.logger.warning(f"清理Word进程时出错：{str(kill_error)}")
                
                # 确保释放锁
                if self._word_lock_file:
                    self._release_word_lock()
                return None
            
            # 其他错误也返回None，避免无限重试
            self.logger.warning(f"DOC文件解析失败：{error_str}")
            if self._word_lock_file:
                self._release_word_lock()
            return None

        finally:
            # === 关键修复4：确保Word进程退出（增强版） ===
            if word:
                try:
                    word.Quit(SaveChanges=False)
                    # 等待进程退出
                    time.sleep(0.5)
                except Exception as e:
                    self.logger.warning(f"关闭Word进程时出错（可忽略）：{str(e)}")
                word = None
            
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            
            # === 关键修复5：强制终止所有Word进程 ===
            time.sleep(0.3)  # 给Word一点时间自己退出
            self._kill_word_processes()
            time.sleep(0.5)  # 再次等待，确保进程完全退出
            
            # === 关键修复6：释放Word COM锁 ===
            self._release_word_lock()
            
            # 检查是否超时
            elapsed = time.time() - start_time
            if elapsed > self.parse_timeout_seconds:
                self.logger.error(f"DOC文件解析超时（{elapsed:.2f}秒 > {self.parse_timeout_seconds}秒）：{file_path}")

    def _convert_doc_to_docx(self, file_path):
        """将.doc转换为.docx再解析（修复版，云端环境兼容）"""
        word = None
        doc = None
        
        try:
            # === 云端环境检查：如果Word COM不可用，返回None ===
            if not self._word_com_available:
                # 使用WARNING级别，因为这是已知的、预期的错误（已在初始化时记录过详细信息）
                self.logger.warning(f"⚠️ Word COM组件不可用，跳过DOC文件转换：{os.path.basename(file_path)}")
                return None
            
            # === 关键修复：转换前清理进程 ===
            self._kill_word_processes()
            time.sleep(0.5)
            
            try:
                pythoncom.CoInitialize()
            except Exception as e:
                self.logger.error(f"COM初始化失败，无法转换：{str(e)}")
                return None

            try:
                word = win32com.client.Dispatch("Word.Application")
                if word is None:
                    raise Exception("无法创建Word.Application对象")
                
                word.Visible = False
                word.DisplayAlerts = 0

                try:
                    # 打开.doc文件（ReadOnly模式）
                    abs_path = os.path.abspath(file_path)
                    self.logger.info(f"尝试转换DOC文件为DOCX：{abs_path}")
                    try:
                        doc = word.Documents.Open(abs_path, ReadOnly=True)
                    except Exception as open_error:
                        error_str = str(open_error)
                        error_code = None
                        
                        # 尝试提取错误代码
                        try:
                            import pywintypes
                            if isinstance(open_error, pywintypes.com_error):
                                error_code = open_error.args[0]
                        except (ImportError, AttributeError):
                            pass
                        
                        # 检查是否是Office检测到文件问题的错误（错误代码 -2147352567 或 -2146821993）
                        is_file_problem = (
                            error_code == -2147352567 or 
                            error_code == -2146821993 or
                            '-2147352567' in error_str or 
                            '-2146821993' in error_str or
                            'Office 检测到此文件存在一个问题' in error_str or 
                            '不能打开此文件' in error_str or
                            '为帮助保护您的计算机' in error_str
                        )
                        
                        if is_file_problem:
                            self.logger.error(f"Word检测到文件存在问题，无法转换（文件可能损坏或包含恶意内容）：{abs_path}")
                            self.logger.error(f"错误代码：{error_code}，错误详情：{error_str}")
                            return None
                        else:
                            raise
                    
                    if doc is None:
                        raise Exception("文档对象为None")

                    # 另存为.docx
                    docx_path = file_path + 'x'
                    doc.SaveAs2(docx_path, FileFormat=12)  # 12 = wdFormatXMLDocument
                    doc.Close(SaveChanges=False)
                    doc = None
                    
                    self.logger.info(f"DOC文件已转换为DOCX：{docx_path}")

                    # 解析转换后的文件
                    result = self._parse_docx(docx_path)

                    # 删除临时文件
                    if os.path.exists(docx_path):
                        try:
                            os.remove(docx_path)
                            self.logger.info(f"已删除临时DOCX文件：{docx_path}")
                        except Exception as e:
                            self.logger.warning(f"删除临时文件失败：{str(e)}")

                    return result

                except Exception as e:
                    self.logger.error(f"DOC转换过程中出错：{str(e)}", exc_info=True)
                    if doc:
                        try:
                            doc.Close(SaveChanges=False)
                        except:
                            pass
                        doc = None
                    raise

                finally:
                    # === 关键修复：强制清理资源 ===
                    if doc:
                        try:
                            doc.Close(SaveChanges=False)
                        except:
                            pass
                        doc = None
                    
                    if word:
                        try:
                            word.Quit(SaveChanges=False)
                        except:
                            pass
                        word = None
                    
                    try:
                        pythoncom.CoUninitialize()
                    except:
                        pass
                    
                    # 强制清理进程
                    time.sleep(0.3)
                    self._kill_word_processes()
                    time.sleep(0.5)
                    
                    # === 关键修复：释放Word COM锁 ===
                    self._release_word_lock()

            except Exception as e:
                self.logger.error(f"Word COM操作失败：{str(e)}", exc_info=True)
                # 确保释放锁
                self._release_word_lock()
                return None

        except Exception as e:
            self.logger.error(f"DOC转换失败 {file_path}：{str(e)}", exc_info=True)
            # 确保释放锁
            self._release_word_lock()
            return None

    def _parse_docx(self, file_path):
        """解析.docx文件（增强表格处理，确保评分表完整提取，支持DOCM格式）"""
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                self.logger.error(f"文件不存在：{file_path}")
                return None
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                self.logger.error(f"文件大小为0：{file_path}")
                return None
            
            # 检查文件大小是否过小
            if file_size < 100:
                self.logger.warning(f"文件过小（{file_size}字节），可能为空或损坏: {file_path}")
                # 检查文件头
                try:
                    with open(file_path, 'rb') as f:
                        header = f.read(8)
                        if header[:2] != b'PK':  # ZIP文件头
                            self.logger.error(f"文件不是有效的DOCX格式（缺少ZIP文件头）: {file_path}")
                            return None
                except Exception as e:
                    self.logger.warning(f"检查文件头失败：{str(e)}")
                    return None
            
            # 尝试打开文档
            try:
                doc = Document(file_path)
            except ValueError as e:
                error_msg = str(e)
                # 检查是否是DOCM格式（启用宏的Word文档）
                if 'macroEnabled' in error_msg or 'application/vnd.ms-word.document.macroEnabled' in error_msg:
                    self.logger.warning(f"文件是DOCM格式（启用宏的Word文档），python-docx无法直接解析：{file_path}")
                    self.logger.info("尝试使用Word COM组件转换为DOCX...")
                    # 如果是DOCM，尝试使用Word COM转换为DOCX
                    if self._word_com_available:
                        result = self._convert_docm_to_docx(file_path)
                        if result and result.strip():
                            return result
                        else:
                            self.logger.warning(f"DOCM转换后内容为空: {file_path}")
                            return None
                    else:
                        # 使用WARNING级别，因为这是已知的、预期的错误（已在初始化时记录过详细信息）
                        self.logger.warning(f"⚠️ Word COM组件不可用，无法解析DOCM文件：{os.path.basename(file_path)}")
                        return None
                elif 'Package not found' in error_msg or 'PackageNotFoundError' in str(type(e).__name__):
                    self.logger.error(f"DOCX文件损坏或格式错误（不是有效的ZIP包）：{file_path}")
                    return None
                else:
                    # 其他ValueError，重新抛出
                    raise
            
            full_text = []

            # === 优化：按文档顺序提取段落和表格，保持原始结构 ===
            # DOCX格式中，段落和表格是交错出现的，需要按顺序提取
            
            # 1. 提取所有元素（段落和表格）按文档顺序
            # 由于python-docx库的限制，我们需要分别提取段落和表格，然后根据位置排序
            # 简化处理：先提取段落，再提取表格（在实际文档中，表格通常在特定位置）
            
            # 提取段落文本
            paragraphs_text = []
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    paragraphs_text.append(para_text)
            
            # 提取表格内容（关键优化：确保所有表格都被提取，添加异常处理）
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_rows = []
                    # 使用try-except保护每一行的处理
                    for row_idx, row in enumerate(table.rows):
                        try:
                            row_cells = []
                            # 尝试获取单元格，如果失败则跳过该行
                            try:
                                cells = row.cells
                            except (IndexError, AttributeError) as cell_error:
                                self.logger.warning(f"表格{table_idx}行{row_idx}访问失败（可能是表格格式问题），跳过：{str(cell_error)[:50]}")
                                continue
                            
                            for cell in cells:
                                try:
                                    cell_text = cell.text.strip()
                                    # 清理单元格文本（移除多余的空白字符，但保留空格）
                                    cell_text = ' '.join(cell_text.split())
                                    row_cells.append(cell_text)
                                except Exception as cell_text_error:
                                    self.logger.debug(f"提取单元格文本失败，使用空字符串：{str(cell_text_error)[:50]}")
                                    row_cells.append("")
                            
                            # 只添加非空行（至少有一个非空单元格）
                            if any(cell.strip() for cell in row_cells):
                                table_rows.append("\t".join(row_cells))
                        except Exception as row_error:
                            # 如果某行访问失败（可能是合并单元格或格式问题），跳过该行
                            self.logger.debug(f"表格{table_idx}行{row_idx}处理失败（可能是有合并单元格），跳过：{str(row_error)[:50]}")
                            continue
                    
                    if table_rows:  # 只保存非空表格
                        tables_data.append({
                            'index': table_idx,
                            'rows': table_rows,
                            'row_count': len(table_rows)
                        })
                        self.logger.debug(f"提取表格 {table_idx + 1}，共 {len(table_rows)} 行")
                except Exception as table_error:
                    # 如果整个表格处理失败，记录警告但继续处理其他表格
                    self.logger.warning(f"表格{table_idx}提取失败（可能是表格格式异常），跳过：{str(table_error)[:100]}")
                    continue
            
            # 2. 合并内容：先添加段落，然后在适当位置插入表格
            # 简化策略：将所有段落放在前面，所有表格放在后面，用标记分隔
            # 这样可以确保所有表格都被保留
            if paragraphs_text:
                full_text.extend(paragraphs_text)
            
            # 3. 添加所有表格（确保评分表不丢失）
            if tables_data:
                full_text.append("\n--- 以下为文档中的表格内容 ---\n")
                for table_info in tables_data:
                    # 添加表格标识
                    full_text.append("[表格开始]")
                    full_text.extend(table_info['rows'])
                    full_text.append("[表格结束]")
                    full_text.append("")  # 表格之间添加空行
                self.logger.info(f"成功提取了 {len(tables_data)} 个表格，共 {sum(t['row_count'] for t in tables_data)} 行")

            result = '\n'.join(full_text)
            text_length = len(result) if result else 0
            self.logger.info(f"DOCX文件解析完成，文本长度：{text_length}字符，段落数：{len(paragraphs_text)}，表格数：{len(tables_data)}")
            
            if not result or not result.strip():
                self.logger.warning(f"DOCX文件解析后内容为空：{file_path}")
                return None
            
            return result

        except ValueError as ve:
            # 处理ValueError（通常是文件格式问题）
            error_msg = str(ve)
            
            # 检查是否是DOCM格式（启用宏的Word文档）
            if 'macroEnabled' in error_msg or 'application/vnd.ms-word.document.macroEnabled' in error_msg:
                self.logger.warning(f"文件是DOCM格式，尝试使用Word COM组件转换：{file_path}")
                if self._word_com_available:
                    result = self._convert_docm_to_docx(file_path)
                    if result and result.strip():
                        return result
                    else:
                        self.logger.warning(f"DOCM转换后内容为空: {file_path}")
                        return None
                else:
                    self.logger.error("Word COM组件不可用，无法解析DOCM文件")
                    return None
            # 检查是否是文件损坏
            elif 'Package not found' in error_msg or 'PackageNotFoundError' in error_msg:
                self.logger.error(f"DOCX文件损坏或格式错误（不是有效的ZIP包）：{file_path}")
                return None
            else:
                # 其他ValueError，记录并返回None
                self.logger.error(f"DOCX文件格式错误 {file_path}：{error_msg}")
                return None
        except Exception as e:
            # 处理其他所有异常
            error_type = type(e).__name__
            error_msg = str(e)
            self.logger.error(f"DOCX解析失败 {file_path}：{error_type}: {error_msg}", exc_info=True)
            return None
    
    def _convert_docm_to_docx(self, file_path):
        """将DOCM文件转换为DOCX再解析（DOCM是启用宏的Word文档）"""
        # DOCM和DOC的处理方式类似，都是使用Word COM组件转换为DOCX
        # 可以重用_convert_doc_to_docx的逻辑，但需要特殊处理
        word = None
        doc = None
        
        try:
            if not self._word_com_available:
                self.logger.error(f"Word COM组件不可用，无法转换DOCM文件：{file_path}")
                return None
            
            # === 关键修复：获取Word COM锁，防止并发访问冲突 ===
            if not self._acquire_word_lock(timeout=60):
                self.logger.warning("Word COM组件正被其他进程使用，无法转换DOCM文件（请稍后重试）")
                return None
            
            try:
                self._kill_word_processes()
                time.sleep(0.5)
                
                try:
                    pythoncom.CoInitialize()
                except Exception as e:
                    self.logger.error(f"COM初始化失败，无法转换DOCM：{str(e)}")
                    self._release_word_lock()
                    return None

                try:
                    word = win32com.client.Dispatch("Word.Application")
                    if word is None:
                        raise Exception("无法创建Word.Application对象")
                    
                    word.Visible = False
                    word.DisplayAlerts = 0

                    try:
                        abs_path = os.path.abspath(file_path)
                        self.logger.info(f"尝试转换DOCM文件为DOCX：{abs_path}")
                        doc = word.Documents.Open(abs_path, ReadOnly=True)
                        
                        if doc is None:
                            raise Exception("文档对象为None")

                        # 另存为.docx（禁用宏）
                        docx_path = file_path.replace('.docm', '.docx') if file_path.endswith('.docm') else file_path + 'x'
                        doc.SaveAs2(docx_path, FileFormat=12)  # 12 = wdFormatXMLDocument
                        doc.Close(SaveChanges=False)
                        doc = None
                        
                        self.logger.info(f"DOCM文件已转换为DOCX：{docx_path}")

                        # 解析转换后的文件
                        result = self._parse_docx(docx_path)

                        # 删除临时文件
                        if os.path.exists(docx_path):
                            try:
                                os.remove(docx_path)
                                self.logger.info(f"已删除临时DOCX文件：{docx_path}")
                            except Exception as e:
                                self.logger.warning(f"删除临时文件失败：{str(e)}")

                        return result

                    except Exception as e:
                        self.logger.error(f"DOCM转换过程中出错：{str(e)}", exc_info=True)
                        if doc:
                            try:
                                doc.Close(SaveChanges=False)
                            except:
                                pass
                            doc = None
                        raise

                    finally:
                        if doc:
                            try:
                                doc.Close(SaveChanges=False)
                            except:
                                pass
                            doc = None
                        
                        if word:
                            try:
                                word.Quit(SaveChanges=False)
                            except:
                                pass
                            word = None
                        
                        try:
                            pythoncom.CoUninitialize()
                        except:
                            pass
                        
                        time.sleep(0.3)
                        self._kill_word_processes()
                        time.sleep(0.5)
                        
                        # === 关键修复：释放Word COM锁 ===
                        self._release_word_lock()

                except Exception as e:
                    self.logger.error(f"Word COM操作失败：{str(e)}", exc_info=True)
                    # 确保释放锁
                    self._release_word_lock()
                    return None

            except Exception as e:
                self.logger.error(f"DOCM转换失败 {file_path}：{str(e)}", exc_info=True)
                # 确保释放锁
                self._release_word_lock()
                return None

        except Exception as e:
            self.logger.error(f"DOCM转换外层异常 {file_path}：{str(e)}", exc_info=True)
            # 确保释放锁
            self._release_word_lock()
            return None

    def _parse_pdf(self, file_path):
        """解析PDF文件（增强版：添加进度和超时控制）"""
        start_time = time.time()
        
        try:
            # 检查文件大小
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                self.logger.warning(f"PDF文件较大（{file_size_mb:.2f}MB），解析可能需要较长时间：{file_path}")
            
            self.logger.info(f"开始解析PDF文件：{file_path}（大小：{file_size_mb:.2f}MB）")
            
            # 普通PDF解析
            if not PYPDF2_AVAILABLE:
                self.logger.error("PyPDF2未安装，无法解析PDF文件。请安装：pip install PyPDF2")
                return None
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                self.logger.info(f"PDF文件共 {total_pages} 页")
                
                text = []
                for i, page in enumerate(reader.pages, 1):
                    # 每10页输出一次进度
                    if i % 10 == 0 or i == total_pages:
                        self.logger.debug(f"PDF解析进度：{i}/{total_pages} 页")
                    
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text.strip())
                    
                    # 检查超时
                    if time.time() - start_time > self.parse_timeout_seconds:
                        self.logger.error(f"PDF解析超时，已解析 {i}/{total_pages} 页")
                        break
                
                elapsed = time.time() - start_time
                self.logger.info(f"PDF文件解析完成，耗时：{elapsed:.2f}秒")
                return '\n'.join(text)

        except Exception as e:
            self.logger.warning(f"普通PDF解析失败，尝试OCR：{str(e)}")
            # 尝试OCR解析扫描件PDF（OCR 很慢，需要更长的超时时间）
            # 检查OCR依赖是否可用
            if not (PIL_AVAILABLE and PYTESSERACT_AVAILABLE and PDF2IMAGE_AVAILABLE):
                self.logger.error("OCR功能不可用：缺少必要的依赖包（PIL/pytesseract/pdf2image）。请安装：pip install pdf2image pytesseract pillow")
                return None
            
            try:
                self.logger.info("开始OCR解析PDF（此过程可能较慢）...")
                ocr_start = time.time()
                
                # 转换PDF为图片（限制页数，避免过慢）
                images = pdf2image.convert_from_path(file_path)
                total_pages = len(images)
                self.logger.info(f"PDF转换为图片完成，共 {total_pages} 页，开始OCR识别...")
                
                text = []
                for i, image in enumerate(images, 1):
                    # 每5页输出一次进度
                    if i % 5 == 0 or i == total_pages:
                        elapsed = time.time() - ocr_start
                        self.logger.info(f"OCR进度：{i}/{total_pages} 页，已耗时：{elapsed:.2f}秒")
                    
                    page_text = pytesseract.image_to_string(image)
                    if page_text:
                        text.append(page_text.strip())
                    
                    # 检查OCR超时
                    if time.time() - ocr_start > self.ocr_timeout_seconds:
                        self.logger.error(f"OCR解析超时，已处理 {i}/{total_pages} 页")
                        break
                
                elapsed = time.time() - ocr_start
                self.logger.info(f"OCR解析完成，耗时：{elapsed:.2f}秒")
                return '\n'.join(text)

            except Exception as ocr_error:
                self.logger.error(f"PDF OCR解析失败：{str(ocr_error)}")
                return None

    def _parse_txt(self, file_path):
        """解析TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read().strip()
        except Exception as e:
            self.logger.error(f"TXT解析失败：{str(e)}")
            return None

    def _parse_excel(self, file_path):
        """解析Excel文件（.xlsx和.xls格式）"""
        try:
            self.logger.info(f"开始解析Excel文件：{file_path}")
            content = []
            
            # 检查文件扩展名
            file_ext = Path(file_path).suffix.lower().lstrip('.')
            
            if file_ext == 'xlsx':
                # 使用openpyxl解析xlsx文件
                from openpyxl import load_workbook
                wb = load_workbook(file_path, data_only=True)
                
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    sheet_content = []
                    sheet_content.append(f"=== 工作表：{sheet_name} ===")
                    
                    # 读取所有行
                    for row in ws.iter_rows(values_only=True):
                        # 过滤空行
                        if any(cell is not None for cell in row):
                            # 将单元格转换为字符串
                            row_str = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                            sheet_content.append(row_str)
                    
                    if len(sheet_content) > 1:  # 跳过空工作表
                        content.extend(sheet_content)
                        content.append('')  # 工作表之间添加空行
                
                wb.close()
            elif file_ext == 'xls':
                # 使用xlrd解析xls文件
                import xlrd
                wb = xlrd.open_workbook(file_path)
                
                for sheet_index in range(wb.nsheets):
                    ws = wb.sheet_by_index(sheet_index)
                    sheet_name = ws.name
                    sheet_content = []
                    sheet_content.append(f"=== 工作表：{sheet_name} ===")
                    
                    # 读取所有行
                    for row_index in range(ws.nrows):
                        row = ws.row_values(row_index)
                        # 过滤空行
                        if any(cell is not None and cell != '' for cell in row):
                            # 将单元格转换为字符串
                            row_str = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                            sheet_content.append(row_str)
                    
                    if len(sheet_content) > 1:  # 跳过空工作表
                        content.extend(sheet_content)
                        content.append('')  # 工作表之间添加空行
                
                wb.close()
            
            if content:
                result = '\n'.join(content)
                self.logger.info(f"Excel文件解析成功，内容长度：{len(result)} 字符")
                return result
            else:
                self.logger.warning(f"Excel文件为空：{file_path}")
                return None
        except Exception as e:
            self.logger.error(f"解析Excel文件失败 {file_path}：{str(e)}")
            import traceback
            self.logger.error(traceback.format_exc())
            return None

    

    def run(self, project_ids=None):
        """批量解析文件（增强版，支持zip文件，添加进程清理）
        
        Args:
            project_ids: 可选，指定要解析的项目ID列表，若为None则解析所有待处理项目
        """
        from utils.db import get_db, TenderProject, update_project, ProjectStatus
        from config import FILES_DIR
        import traceback  # 新增

        # === 关键修复：开始前清理所有Word进程 ===
        self._kill_word_processes()

        db = next(get_db())
        from sqlalchemy import or_
        # 构建查询
        # 排除已经标记为多次失败的项目（error_msg中包含"[跳过-多次失败]"标记）
        query = db.query(TenderProject).filter(
            TenderProject.status.in_([ProjectStatus.DOWNLOADED, ProjectStatus.ERROR])
        ).filter(
            or_(
                TenderProject.error_msg.is_(None),
                ~TenderProject.error_msg.like('%[跳过-多次失败]%')
            )
        )
        
        # 如果指定了项目ID，则只处理这些项目
        if project_ids and len(project_ids) > 0:
            query = query.filter(TenderProject.id.in_(project_ids))
        
        projects = query.all()

        self.logger.info(f"待解析项目数：{len(projects)}")
        
        total_start_time = time.time()
        processed_count = 0
        success_count = 0
        error_count = 0

        for idx, project in enumerate(projects, 1):
            try:
                self.logger.info(f"[{idx}/{len(projects)}] 开始解析项目：{project.project_name}（ID：{project.id}）")
                processed_count += 1

                # === 关键修复：每2个文件清理一次进程 ===
                if idx > 1 and idx % 2 == 0:
                    self.logger.info(f"清理Word进程（已处理 {idx} 个文件）...")
                    self._kill_word_processes()
                    time.sleep(0.5)

                # 检查文件路径
                file_path = project.file_path
                if not file_path:
                    update_project(db, project.id, {
                        "status": ProjectStatus.ERROR,
                        "error_msg": "文件路径为空，可能是下载失败"
                    })
                    self.logger.warning(f"跳过项目 {project.project_name}：文件路径为空")
                    continue

                # 处理相对路径
                if not os.path.isabs(file_path):
                    file_path = os.path.join(FILES_DIR, file_path)

                # 检查文件是否存在
                if not os.path.exists(file_path):
                    update_project(db, project.id, {
                        "status": ProjectStatus.ERROR,
                        "error_msg": f"文件不存在：{file_path}"
                    })
                    self.logger.warning(f"跳过项目 {project.project_name}：文件不存在")
                    error_count += 1
                    continue
                
                # 检查文件大小，小文件可能是空文件或损坏文件
                file_size = os.path.getsize(file_path)
                # 降低阈值到2KB，并添加文件头检查，避免误判有效文件
                if file_size < 2048:  # 小于2KB的文件，很可能是空文件或损坏文件
                    # 对于非常小的文件，检查文件头是否有效
                    is_valid_file = False
                    try:
                        with open(file_path, 'rb') as f:
                            header = f.read(8)
                            # 检查是否是有效的Word文档（OLE2格式）
                            if header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                                is_valid_file = True
                            # 检查是否是有效的DOCX/ZIP格式
                            elif header[:2] == b'PK':
                                is_valid_file = True
                            # 检查是否是PDF格式
                            elif header[:4] == b'%PDF':
                                is_valid_file = True
                    except:
                        pass
                    
                    if not is_valid_file:
                        update_project(db, project.id, {
                            "status": ProjectStatus.ERROR,
                            "error_msg": f"文件过小（{file_size}字节），可能是空文件或损坏文件"
                        })
                        self.logger.warning(f"跳过项目 {project.project_name}：文件过小（{file_size}字节）且文件头无效")
                        error_count += 1
                        continue
                    else:
                        self.logger.info(f"文件较小（{file_size}字节），但文件头有效，继续解析：{project.project_name}")
                
                # 检查文件扩展名
                file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')
                if file_ext not in self.supported_formats:
                    # 检查是否是压缩文件
                    if file_ext not in self.archive_formats:
                        update_project(db, project.id, {
                            "status": ProjectStatus.ERROR,
                            "error_msg": f"不支持的文件格式：{file_ext}"
                        })
                        self.logger.warning(f"跳过项目 {project.project_name}：不支持的文件格式 {file_ext}")
                        error_count += 1
                        continue
                
                # 解析文件（添加超时保护，使用线程实现真正的超时机制）
                self.logger.info(f"开始解析文件：{file_path}（大小：{file_size}字节）")
                parse_start_time = time.time()
                
                # 使用超时机制，防止单个文件解析卡住
                content = None
                parse_error = None
                timeout_occurred = False
                
                def parse_with_timeout():
                    """在单独线程中执行解析，支持超时中断"""
                    nonlocal content, parse_error
                    try:
                        content = self.parse_file(file_path, project.id)
                    except Exception as e:
                        parse_error = e
                
                # 创建解析线程
                parse_thread = threading.Thread(target=parse_with_timeout, daemon=True)
                parse_thread.start()
                
                # 等待解析完成或超时
                parse_thread.join(timeout=self.parse_timeout_seconds)
                
                parse_elapsed = time.time() - parse_start_time
                
                # 检查是否超时
                if parse_thread.is_alive():
                    timeout_occurred = True
                    self.logger.error(f"⏱️ 文件解析超时（超过{self.parse_timeout_seconds}秒），强制中断：{file_path}")
                    # 清理可能的残留进程
                    self._kill_word_processes()
                    # 等待一小段时间让线程有机会退出
                    time.sleep(1)
                    content = None
                    parse_error = TimeoutError(f"文件解析超时（超过{self.parse_timeout_seconds}秒）")
                elif parse_error:
                    self.logger.error(f"文件解析异常，耗时 {parse_elapsed:.2f}秒：{str(parse_error)}")
                    # 清理可能的残留进程
                    self._kill_word_processes()
                    content = None
                elif parse_elapsed > 300:
                    # 如果解析时间超过5分钟，记录警告
                    self.logger.warning(f"⚠️ 文件解析耗时较长：{parse_elapsed:.2f}秒，文件：{file_path}")
                
                # 如果超时，更新错误信息并继续处理下一个文件
                if timeout_occurred:
                    # 检查失败次数
                    parse_fail_count = 0
                    import re
                    if project.error_msg:
                        match = re.search(r'\[解析失败(\d+)次\]', project.error_msg)
                        if match:
                            parse_fail_count = int(match.group(1)) + 1
                        else:
                            parse_fail_count = 1
                    else:
                        parse_fail_count = 1
                    
                    error_msg = f"文件解析超时（超过{self.parse_timeout_seconds}秒）"
                    if parse_fail_count >= 3:
                        error_msg = f"{error_msg} [解析失败{parse_fail_count}次] [跳过-多次失败]"
                        self.logger.warning(f"⚠️ 项目 {project.project_name}（ID：{project.id}）已失败{parse_fail_count}次，标记为跳过")
                        update_project(db, project.id, {
                            "status": ProjectStatus.ERROR,
                            "error_msg": error_msg
                        })
                    else:
                        error_msg = f"{error_msg} [解析失败{parse_fail_count}次]"
                        # 自动重试：重置状态为DOWNLOADED，让它重新进入解析流程
                        self.logger.info(f"🔄 项目 {project.project_name}（ID：{project.id}）解析失败第{parse_fail_count}次，自动重置状态准备重试")
                        update_project(db, project.id, {
                            "status": ProjectStatus.DOWNLOADED,  # 重置为DOWNLOADED状态，下次解析时会重新处理
                            "error_msg": error_msg,
                            "evaluation_content": None  # 清空之前可能的部分解析内容
                        })
                    error_count += 1
                    self.logger.error(f"❌ 解析失败：{project.project_name}（{error_msg}）")
                    continue  # 跳过当前文件，继续处理下一个
                
                # 如果解析异常（非超时），也更新错误信息
                if parse_error and not timeout_occurred:
                    # 检查失败次数
                    parse_fail_count = 0
                    import re
                    if project.error_msg:
                        match = re.search(r'\[解析失败(\d+)次\]', project.error_msg)
                        if match:
                            parse_fail_count = int(match.group(1)) + 1
                        else:
                            base_error = re.sub(r'\[解析失败\d+次\].*', '', project.error_msg).strip()
                            current_base_error = re.sub(r'\[解析失败\d+次\].*', '', str(parse_error)).strip()
                            if base_error == current_base_error or current_base_error in base_error:
                                parse_fail_count = 2
                            else:
                                parse_fail_count = 1
                    else:
                        parse_fail_count = 1
                    
                    error_msg = f"解析异常：{str(parse_error)[:200]}"
                    if parse_fail_count >= 3:
                        error_msg = f"{error_msg} [解析失败{parse_fail_count}次] [跳过-多次失败]"
                        self.logger.warning(f"⚠️ 项目 {project.project_name}（ID：{project.id}）已失败{parse_fail_count}次，标记为跳过")
                        update_project(db, project.id, {
                            "status": ProjectStatus.ERROR,
                            "error_msg": error_msg
                        })
                    else:
                        error_msg = f"{error_msg} [解析失败{parse_fail_count}次]"
                        # 自动重试：重置状态为DOWNLOADED，让它重新进入解析流程
                        self.logger.info(f"🔄 项目 {project.project_name}（ID：{project.id}）解析失败第{parse_fail_count}次，自动重置状态准备重试")
                        update_project(db, project.id, {
                            "status": ProjectStatus.DOWNLOADED,  # 重置为DOWNLOADED状态，下次解析时会重新处理
                            "error_msg": error_msg,
                            "evaluation_content": None  # 清空之前可能的部分解析内容
                        })
                    error_count += 1
                    self.logger.error(f"❌ 解析失败：{project.project_name}（{error_msg}）")
                    continue  # 跳过当前文件，继续处理下一个
                
                # 详细记录解析结果（只有在没有超时和异常的情况下才处理）
                if not timeout_occurred and not parse_error:
                    if content:
                        content_length = len(content) if content else 0
                        self.logger.info(f"解析成功，内容长度：{content_length}字符")
                    # 修复字段名错误（evaluation_content而非content）
                    update_project(db, project.id, {
                        "evaluation_content": content,
                        "status": ProjectStatus.PARSED
                    })
                    success_count += 1
                    elapsed = time.time() - total_start_time
                    avg_time = elapsed / processed_count if processed_count > 0 else 0
                    remaining = len(projects) - processed_count
                    estimated_remaining_time = avg_time * remaining if remaining > 0 else 0
                    self.logger.info(f"✅ 解析成功：{project.project_name}（成功：{success_count}，失败：{error_count}，预计剩余：{estimated_remaining_time:.0f}秒）")
                else:
                    self.logger.error(f"❌ 解析失败：{project.project_name}（内容为空），文件路径：{file_path}")
                    # 检查文件是否存在
                    if not os.path.exists(file_path):
                        error_msg = f"文件不存在：{file_path}"
                    else:
                        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
                        error_msg = f"解析内容为空（文件大小：{file_size}字节）"
                    
                    # 检查失败次数：通过error_msg中的失败计数来判断
                    parse_fail_count = 0
                    import re
                    if project.error_msg:
                        # 检查error_msg中是否包含失败计数标记
                        match = re.search(r'\[解析失败(\d+)次\]', project.error_msg)
                        if match:
                            parse_fail_count = int(match.group(1)) + 1  # 增加失败次数
                        else:
                            # 如果没有失败计数，检查是否是相同类型的错误
                            # 提取错误类型（去掉失败次数标记）
                            base_error = re.sub(r'\[解析失败\d+次\].*', '', project.error_msg).strip()
                            current_base_error = re.sub(r'\[解析失败\d+次\].*', '', error_msg).strip()
                            if base_error == current_base_error or current_base_error in base_error:
                                parse_fail_count = 2  # 相同错误，设为2次（下次就是3次）
                            else:
                                parse_fail_count = 1  # 不同错误，重新计数
                    else:
                        parse_fail_count = 1
                    
                    # 如果失败次数达到3次，标记为跳过
                    if parse_fail_count >= 3:
                        error_msg = f"{error_msg} [解析失败{parse_fail_count}次] [跳过-多次失败]"
                        self.logger.warning(f"⚠️ 项目 {project.project_name}（ID：{project.id}）已失败{parse_fail_count}次，标记为跳过，不再尝试解析")
                        update_project(db, project.id, {
                            "status": ProjectStatus.ERROR,
                            "error_msg": error_msg
                        })
                    else:
                        error_msg = f"{error_msg} [解析失败{parse_fail_count}次]"
                        # 自动重试：重置状态为DOWNLOADED，让它重新进入解析流程
                        self.logger.info(f"🔄 项目 {project.project_name}（ID：{project.id}）解析失败第{parse_fail_count}次，自动重置状态准备重试")
                        update_project(db, project.id, {
                            "status": ProjectStatus.DOWNLOADED,  # 重置为DOWNLOADED状态，下次解析时会重新处理
                            "error_msg": error_msg,
                            "evaluation_content": None  # 清空之前可能的部分解析内容
                        })
                    error_count += 1
                    self.logger.error(f"❌ 解析失败：{project.project_name}（{error_msg}）")

            except Exception as e:
                error_count += 1
                base_error_msg = f"{str(e)} \n {traceback.format_exc()[:500]}"  # 增加堆栈信息
                
                # 检查失败次数：通过error_msg中的失败计数来判断
                parse_fail_count = 0
                if project.error_msg:
                    # 检查error_msg中是否包含失败计数标记
                    import re
                    match = re.search(r'\[解析失败(\d+)次\]', project.error_msg)
                    if match:
                        parse_fail_count = int(match.group(1))
                    # 如果error_msg相似（包含相同的错误类型），说明是重复失败
                    if project.error_msg and (str(e) in project.error_msg or project.error_msg in str(e)):
                        parse_fail_count += 1
                    else:
                        parse_fail_count = 1
                else:
                    parse_fail_count = 1
                
                # 如果失败次数达到3次，标记为跳过
                if parse_fail_count >= 3:
                    error_msg = f"{base_error_msg} [解析失败{parse_fail_count}次] [跳过-多次失败]"
                    self.logger.warning(f"⚠️ 项目 {project.project_name}（ID：{project.id}）已失败{parse_fail_count}次，标记为跳过，不再尝试解析")
                    update_project(db, project.id, {
                        "status": ProjectStatus.ERROR,
                        "error_msg": error_msg
                    })
                else:
                    error_msg = f"{base_error_msg} [解析失败{parse_fail_count}次]"
                    # 自动重试：重置状态为DOWNLOADED，让它重新进入解析流程
                    self.logger.info(f"🔄 项目 {project.project_name}（ID：{project.id}）解析失败第{parse_fail_count}次，自动重置状态准备重试")
                    update_project(db, project.id, {
                        "status": ProjectStatus.DOWNLOADED,  # 重置为DOWNLOADED状态，下次解析时会重新处理
                        "error_msg": error_msg,
                        "evaluation_content": None  # 清空之前可能的部分解析内容
                    })
                self.logger.error(f"❌ 处理项目失败 {project.project_name}：{str(e)}")
                
                # === 关键修复：出错时也清理进程 ===
                self._kill_word_processes()
                time.sleep(0.5)
                continue

        # === 关键修复：最后清理一次 ===
        self._kill_word_processes()

        db.close()
        total_elapsed = time.time() - total_start_time
        self.logger.info("=" * 60)
        self.logger.info(f"文件解析完成！")
        self.logger.info(f"总计：{processed_count} 个，成功：{success_count} 个，失败：{error_count} 个")
        self.logger.info(f"总耗时：{total_elapsed:.2f}秒，平均：{total_elapsed/processed_count if processed_count > 0 else 0:.2f}秒/个")
        self.logger.info("=" * 60)